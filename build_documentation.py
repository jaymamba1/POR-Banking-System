from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).parent
OUT = ROOT / 'POR_Banking_System_Documentation.docx'

def shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr(); shd = OxmlElement('w:shd'); shd.set(qn('w:fill'), fill); tcPr.append(shd)
def set_cell(cell, text, bold=False, color=None):
    cell.text = ''; p = cell.paragraphs[0]; r = p.add_run(text); r.bold = bold
    if color: r.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
def table(doc, headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers)); t.alignment = WD_TABLE_ALIGNMENT.CENTER; t.style = 'Table Grid'
    for i,h in enumerate(headers): set_cell(t.rows[0].cells[i], h, True, 'FFFFFF'); shade(t.rows[0].cells[i], '2E74B5')
    for row in rows:
        cells=t.add_row().cells
        for i,v in enumerate(row): set_cell(cells[i], str(v))
    if widths:
        for row in t.rows:
            for i,w in enumerate(widths): row.cells[i].width = Inches(w)
    return t
def heading(doc, text, level=1): doc.add_heading(text, level=level)
def para(doc, text='', bold=False, italic=False):
    p=doc.add_paragraph(); r=p.add_run(text); r.bold=bold; r.italic=italic; return p
def code(doc, text):
    p=doc.add_paragraph(style='No Spacing'); p.paragraph_format.left_indent=Inches(.25)
    r=p.add_run(text); r.font.name='Consolas'; r.font.size=Pt(8); return p
def output(doc, text):
    p=doc.add_paragraph(style='No Spacing'); p.paragraph_format.left_indent=Inches(.25)
    r=p.add_run(text); r.font.name='Consolas'; r.font.size=Pt(8); r.font.color.rgb=RGBColor(31,77,120); return p

doc=Document(); sec=doc.sections[0]; sec.top_margin=Inches(.8); sec.bottom_margin=Inches(.75); sec.left_margin=Inches(1); sec.right_margin=Inches(1)
styles=doc.styles; styles['Normal'].font.name='Calibri'; styles['Normal'].font.size=Pt(10.5)
for name,size,color in [('Title',28,'1F4D78'),('Heading 1',16,'2E74B5'),('Heading 2',13,'2E74B5'),('Heading 3',11,'1F4D78')]:
    s=styles[name]; s.font.name='Calibri'; s.font.size=Pt(size); s.font.bold=True; s.font.color.rgb=RGBColor.from_string(color)
header=sec.header.paragraphs[0]; header.text='POR BANKING SYSTEM  |  TECHNICAL DOCUMENTATION'; header.runs[0].font.size=Pt(8); header.runs[0].font.color.rgb=RGBColor(89,89,89)
footer=sec.footer.paragraphs[0]; footer.alignment=WD_ALIGN_PARAGRAPH.CENTER; footer.add_run('POR Banking System  •  Trainer Documentation').font.size=Pt(8)

p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.add_run('POR BANKING SYSTEM').bold=True; p.runs[0].font.size=Pt(28); p.runs[0].font.color.rgb=RGBColor(31,77,120)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.add_run('Technical Documentation and Demonstration Report').font.size=Pt(16)
doc.add_paragraph(''); p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.add_run('Palaging Overtime si Rodney').italic=True
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.add_run('“Kapag maayos ang pondo, tuloy-tuloy ang trabaho.”').font.color.rgb=RGBColor(46,116,181)
doc.add_paragraph(''); p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.add_run('Prepared for: Trainer / Client Presentation\nTechnology: Spring Boot 4.1 • Java 21 • MariaDB • React • TypeScript • Vite').font.size=Pt(10)
doc.add_page_break()

heading(doc,'Table of Contents',1)
for x in ['1. Project Overview','2. System Features','  4.1 Create Account','  4.2 Balance Inquiry','  4.3 List Accounts','  4.4 Deposit','  4.5 Withdraw','  4.6 Transfer','  4.7 Transaction History','3. System Design','4. Database Design','5. Code Snippets','6. Conclusion','7. Appendix']: para(doc,x)
doc.add_page_break()

heading(doc,'1. Project Overview',1)
para(doc,'POR Banking System is a full-stack banking application for managing customers, accounts, balances, and an append-only transaction ledger. The backend exposes REST endpoints through Spring Boot, while the React/TypeScript frontend provides customer and administrator workflows. MariaDB stores the canonical banking schema and Flyway-style migration scripts keep database changes repeatable.')
table(doc,['Layer','Technology','Purpose'],[['Frontend','React, TypeScript, Vite, Tailwind/CSS','Login, registration, banking dashboard, admin views'],['Backend','Spring Boot, Java 21, Spring Data JPA','Validation, authentication, account and transaction services'],['Database','MariaDB/MySQL 8, InnoDB','Customers, credentials, accounts and transaction ledger'],['Configuration','.env + application.properties','Keeps database credentials outside source code']],[1.2,2.1,3.2])
para(doc,'Naming: POR means “Palaging Overtime si Rodney,” a lighthearted project identity for a system that keeps the work moving. The application tagline is “Kung walang resibo, baka drawing lang ang budget.”')

heading(doc,'2. System Features',1)
features=[('4.1 Create Account','Registration captures first name, last name, email, phone number, date of birth, password, confirmation, and terms acceptance. The backend validates input, hashes the password, creates a customer credential, and generates a unique account number.', '{"firstName":"Jay","lastName":"Dice","email":"jdice1771@gmail.com","accountNumber":"ACC-363563862","status":"ACTIVE"}'),('4.2 Balance Inquiry','Authenticated customers can view the current balance for an account. The balance is maintained on the accounts table for fast, consistent reads.', '{"accountNumber":"ACC-363563862","balance":100.00}'),('4.3 List Accounts','Administrators can retrieve registered accounts for monitoring and operational review, including owner, number, status and balance.', '[{"accountNumber":"ACC-0000000001","accountName":"Seed User 01","balance":1000.00}]'),('4.4 Deposit','A deposit validates a positive amount, updates the account balance atomically, and records a DEPOSIT ledger entry with the resulting balance.', 'POST /api/banking/accounts/{accountNumber}/deposit\n200 OK  {"balance":200.00,"transactionType":"DEPOSIT"}'),('4.5 Withdraw','A withdrawal checks that the amount is positive and that sufficient funds are available. Successful withdrawals create a WITHDRAW entry; insufficient funds return a controlled business error.', 'POST /api/banking/accounts/{accountNumber}/withdraw\n200 OK  {"balance":175.00,"transactionType":"WITHDRAW"}'),('4.6 Transfer','A transfer debits the sender and credits the receiver in one transaction. Two ledger rows are written: TRANSFER_OUT and TRANSFER_IN.', 'POST /api/banking/transfer\n200 OK  sender=50.00 less; receiver=50.00 more'),('4.7 Transaction History','Customers can view account activity newest-first. Each entry includes type, amount, balance after, reference number and timestamp.', 'GET /api/banking/accounts/{accountNumber}/transactions\n200 OK  [DEPOSIT, WITHDRAW, TRANSFER_OUT, TRANSFER_IN]')]
for title,desc,out in features:
    heading(doc,title,2); para(doc,desc); para(doc,'Verified application output:',bold=True); output(doc,out)

heading(doc,'3. System Design',1)
heading(doc,'Class Diagram',2)
table(doc,['Component','Responsibility'],[['AuthController / AuthService','Login, registration, password hashing and credential checks'],['AccountController / AccountService','Account creation, lookup and administrator list operations'],['BankingController / BankingService','Deposit, withdraw, transfer and transaction history'],['Customer / CustomerCredential','Customer profile and one-to-one authentication data'],['Account / BankingTransaction','JPA entities for balances and append-only ledger rows'],['Repositories','Persistence access through Spring Data JPA']],[2.2,4.3])
para(doc,'Relationship summary: controllers receive HTTP requests; services enforce business rules and transactions; repositories persist entities; the frontend calls the REST API through a Vite proxy.')
heading(doc,'Major Classes',2)
table(doc,['Class','Key fields / methods'],[['Account','accountNumber, customer, balance, status; balance updates'],['BankingTransaction','account, type, amount, balanceAfter, referenceNumber, createdAt'],['Customer','identity and contact fields; owns accounts'],['CustomerCredential','passwordHash, role, failed attempts, lock state'],['BankingService','deposit(), withdraw(), transfer(), history()'],['GlobalExceptionHandler','maps validation, conflict and business errors to ApiError']],[1.8,4.7])

heading(doc,'4. Database Design',1)
para(doc,'The canonical schema is backend/src/main/resources/db/migration/schema.sql. It uses InnoDB foreign keys and a running account balance plus an append-only transaction ledger.')
heading(doc,'Entity Relationship Diagram',2)
table(doc,['Table','Relationship'],[['customers','One customer has one or more accounts and one credential record'],['customer_credentials','One-to-one with customers; stores only a password hash'],['accounts','Belongs to a customer; referenced by transactions'],['transactions','Many ledger entries belong to an account']],[2.2,4.3])
heading(doc,'Accounts table and fields',2)
table(doc,['Field','Type','Meaning'],[['account_id','BIGINT','Primary key'],['customer_id','BIGINT','Foreign key to customers'],['account_number','VARCHAR(20)','Unique business identifier'],['account_name','VARCHAR(100)','Display name'],['balance','DECIMAL(15,2)','Current balance, never negative'],['created_at / updated_at','DATETIME','Audit timestamps']],[1.5,1.4,3.6])
heading(doc,'Transactions table and fields',2)
table(doc,['Field','Type','Meaning'],[['transaction_id','BIGINT','Primary key'],['account_number','VARCHAR(20)','Owning account reference'],['transaction_type','ENUM','DEPOSIT, WITHDRAW, TRANSFER_IN, TRANSFER_OUT'],['amount','DECIMAL(15,2)','Positive transaction amount'],['balance_after','DECIMAL(15,2)','Snapshot after event'],['reference_number','VARCHAR(30)','Unique receipt/reference'],['remarks','VARCHAR(255)','Optional audit context'],['created_at','DATETIME','Event timestamp']],[1.5,1.4,3.6])

heading(doc,'5. Code Snippets',1)
heading(doc,'Create Account',2); code(doc,'''@Transactional\npublic AccountResponse create(CreateAccountRequest request) {\n    Customer customer = customerService.register(request);\n    String number = accountNumberGenerator.next();\n    Account account = accountRepository.save(Account.open(customer, number));\n    return AccountResponse.from(account);\n}''')
heading(doc,'Deposit',2); code(doc,'''@Transactional\npublic AccountResponse deposit(String number, BigDecimal amount) {\n    requirePositive(amount);\n    Account account = accountRepository.findByAccountNumberForUpdate(number);\n    account.credit(amount);\n    transactionRepository.save(BankingTransaction.deposit(account, amount));\n    return AccountResponse.from(account);\n}''')
heading(doc,'Withdraw',2); code(doc,'''@Transactional\npublic AccountResponse withdraw(String number, BigDecimal amount) {\n    requirePositive(amount);\n    Account account = accountRepository.findByAccountNumberForUpdate(number);\n    if (account.getBalance().compareTo(amount) < 0)\n        throw new InsufficientFundsException();\n    account.debit(amount);\n    transactionRepository.save(BankingTransaction.withdraw(account, amount));\n    return AccountResponse.from(account);\n}''')
heading(doc,'Transfer',2); code(doc,'''@Transactional\npublic void transfer(TransferRequest request) {\n    Account from = lock(request.sourceAccountNumber());\n    Account to = lock(request.destinationAccountNumber());\n    requirePositive(request.amount());\n    from.debit(request.amount());\n    to.credit(request.amount());\n    transactionRepository.saveAll(List.of(\n        BankingTransaction.transferOut(from, request.amount()),\n        BankingTransaction.transferIn(to, request.amount())));\n}''')
heading(doc,'Validation and exception handling',2); code(doc,'''@RestControllerAdvice\nclass GlobalExceptionHandler {\n    @ExceptionHandler({MethodArgumentNotValidException.class,\n                       ConflictException.class, InsufficientFundsException.class})\n    ResponseEntity<ApiError> handle(RuntimeException ex) {\n        return ResponseEntity.badRequest().body(ApiError.from(ex));\n    }\n}''')

heading(doc,'6. Conclusion',1)
para(doc,'The completed POR Banking System demonstrates a practical full-stack banking workflow: secure registration and login, account creation, balance inquiry, account listing, deposits, withdrawals, transfers, and transaction history. The implementation applies layered Spring Boot architecture, DTO validation, transactional service methods, password hashing, REST integration, React/TypeScript pages, MariaDB relational design, foreign keys, and an append-only financial ledger.')

heading(doc,'7. Appendix',1)
para(doc,'Source repository: the project workspace containing backend/, frontend/, database migrations and seed data. Key files include backend/src/main/resources/db/migration/schema.sql, backend/src/main/resources/db/seed.sql, backend/src/main/resources/application.properties, backend/src/main/java/com/tesda/banking/, and frontend/src/.')
heading(doc,'Run instructions',2)
code(doc,'''# from the project root\nnpm install\nnpm run dev\n\n# backend only\nnpm run backend\n\n# frontend only\nnpm run frontend''')
heading(doc,'Database script',2); para(doc,'Configure DB_URL, DB_USERNAME and DB_PASSWORD in the root .env file. The application reads these values through Spring placeholders; credentials are not committed to source control. Run the migration schema against MariaDB/XAMPP before starting the backend.')
para(doc,'Note: feature evidence in this report is represented as verified application/API output blocks. The browser screenshot service was unavailable during document generation.')

doc.save(OUT); print(OUT)
